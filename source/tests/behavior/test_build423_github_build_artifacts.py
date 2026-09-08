"""Build 423 regression coverage for GitHub Build Process separation."""
from __future__ import annotations

from pathlib import Path

from docx import Document
import pytest

pytestmark = pytest.mark.behavior
__version__ = "v446"

ROOT = Path(__file__).resolve().parents[2]


def _doc_text(path: Path) -> str:
    d = Document(path)
    chunks = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_build423_public_version_and_current_documents():
    import tlo_version as V

    assert V.VERSION == "v446"
    assert V.PUBLIC_VERSION == "1.6"
    assert V.BUNDLE_BUILD == 446
    assert V.DISPLAY_VERSION == "v1.6 Build 446"
    assert (ROOT / "TLO_Inventory_Requirements_Working_v446.docx").is_file()
    assert (ROOT / "TLO_Inventory_User_Manual_v446.rtf").is_file()


def test_build423_source_bundle_contains_no_github_build_process_artifacts():
    forbidden = [
        ROOT / "Run-TLO-GitHub-Build.ps1",
        ROOT / "Create-TLOArtifactSigningMetadata.ps1",
    ]
    assert not any(path.exists() for path in forbidden)
    assert not list(ROOT.glob("TLO_GitHub_Build_Process_Requirements_v*.docx"))
    assert not list(ROOT.glob("TLO_GitHub_Build_Process_v*.zip"))


def test_build423_tlo_requirements_record_strict_separation_rule():
    req = ROOT / "TLO_Inventory_Requirements_Working_v446.docx"
    text = _doc_text(req)
    assert "Current document version: v446 (v1.6 Build 446)." in text
    assert "source bundle and the independently versioned GitHub Build Process are separate artifacts" in text
    assert "shall contain no GitHub Build Process files" in text
    assert "Run-TLO-GitHub-Build.ps1" in text
    assert "Create-TLOArtifactSigningMetadata.ps1" in text
    assert "distributed only as its own independently versioned package" in text
