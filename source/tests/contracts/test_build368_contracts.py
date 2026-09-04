"""Build 375 simple deleteDupes mismatch logging contracts."""

__version__ = "v433"

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

pytestmark = pytest.mark.contract
ROOT = Path(__file__).resolve().parents[2]


def _docx_text(name: str) -> str:
    doc = Document(ROOT / name)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_build368_source_records_simple_mismatch_reasons_separately():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "deleteDupesMismatches.txt" in source
    assert 'return "different number of files"' in source
    assert 'return "different sub-structure"' in source
    assert 'return "different file names"' in source
    assert 'different sizes"' in source
    assert "csv.writer" in source


def test_build368_requirements_document_simple_mismatch_log():
    text = _docx_text("TLO_Inventory_Requirements_Working_v433.docx")
    assert "deleteDupesMismatches.txt" in text
    assert "folder name, folder name, simple reason" in text
    assert "different number of files" in text
    assert "different sub-structure" in text
    assert "different sizes" in text


def test_build368_manual_documents_simple_mismatch_log():
    text = (ROOT / "TLO_Inventory_User_Manual_v433.rtf").read_text(encoding="utf-8", errors="ignore")
    assert "deleteDupesMismatches.txt" in text
    assert "different number of files" in text
    assert "different sub-structure" in text
    assert "different sizes" in text


def test_build368_archives_build367_change_log():
    with ZipFile(ROOT / "old-change-logs.zip") as zf:
        assert "CHANGES_v367.txt" in zf.namelist()
