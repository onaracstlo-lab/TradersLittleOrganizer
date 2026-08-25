"""Build 375 copy-to-copy duplicate cluster contracts."""

__version__ = "v397"

from pathlib import Path

import pytest

pytestmark = pytest.mark.contract
from zipfile import ZipFile

from docx import Document

ROOT = Path(__file__).resolve().parents[2]


def _docx_text(name: str) -> str:
    doc = Document(ROOT / name)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_build370_source_compares_copy_family_members_with_each_other():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "Multiple numbered copies therefore compare" in source
    assert "manifest_clusters" in source
    assert "_keeper_preference_key" in source
    assert "copies are compared with one another" in source


def test_build370_requirements_include_x_copy2_copy3_example():
    text = _docx_text("TLO_Inventory_Requirements_Working_v397.docx")
    assert "X, X (copy2), and X (copy3)" in text
    assert "copy2 and copy3 match one another" in text
    assert "X and copy2 remain while copy3 is relocated" in text


def test_build370_manual_and_faq_explain_cluster_preference():
    manual = (ROOT / "TLO_Inventory_User_Manual_v397.rtf").read_text(encoding="utf-8", errors="ignore")
    faq = (ROOT / "TLO-FAQ.txt").read_text(encoding="utf-8")
    for text in (manual, faq):
        assert "copy2" in text
        assert "copy3" in text
        assert "unsuffixed" in text


def test_build370_archives_build369_change_log():
    with ZipFile(ROOT / "old-change-logs.zip") as zf:
        assert "CHANGES_v369.txt" in zf.namelist()
