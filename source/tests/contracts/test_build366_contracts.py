"""Current duplicate-cleanup discovery/keeper contracts (originated in Build 366)."""

__version__ = "v406"

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

pytestmark = pytest.mark.contract
ROOT = Path(__file__).resolve().parents[2]


def _docx_text(name: str) -> str:
    doc = Document(ROOT / name)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_duplicate_source_uses_content_equivalence_clusters_and_preferred_keeper():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "_candidate_components_for_parent" in source
    assert "_keeper_preference_key" in source
    assert "content-equivalence cluster" in source
    assert "manifest_clusters" in source
    assert "lowest copy number is preferred" in source


def test_requirements_define_content_cluster_keeper_rule():
    text = _docx_text("TLO_Inventory_Requirements_Working_v406.docx")
    assert "18.2 Candidate Discovery and Content-Equivalence Clusters" in text
    assert "numbered copies shall be compared with one another" in text
    assert "Prefer an unsuffixed folder as the keeper" in text
    assert "keep the lowest-numbered copy" in text
    assert "X and copy2 remain while copy3 is relocated" in text


def test_manual_documents_content_cluster_keeper_rule():
    text = (ROOT / "TLO_Inventory_User_Manual_v406.rtf").read_text(encoding="utf-8", errors="ignore")
    assert "Copies are compared with one another as well as with unsuffixed folders" in text
    assert "content-equivalence clusters" in text
    assert "lowest-numbered copy is kept" in text
    assert "TLO keeps X and copy2 and moves copy3" in text


def test_faq_mentions_copy_to_copy_comparison():
    text = (ROOT / "TLO-FAQ.txt").read_text(encoding="utf-8")
    assert "compares copies with one another" in text
    assert "content-equivalence cluster" in text
    assert "keeps X and copy2 and moves copy3" in text


def test_build366_archives_build365_change_log():
    with ZipFile(ROOT / "old-change-logs.zip") as zf:
        assert "CHANGES_v365.txt" in zf.namelist()
