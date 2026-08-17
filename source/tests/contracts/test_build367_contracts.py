"""Build 373 partition-root duplicate holding-folder contracts."""

__version__ = "v373"

from pathlib import Path
from zipfile import ZipFile

import pytest
from docx import Document

pytestmark = pytest.mark.contract
ROOT = Path(__file__).resolve().parents[2]


def _docx_text(name: str) -> str:
    doc = Document(ROOT / name)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.extend(p.text for p in cell.paragraphs)
    return "\n".join(chunks)


def test_build367_source_uses_partition_root_duplicates_holding_folder():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "_partition_root_for_path" in source
    assert "_duplicates_root_for_search" in source
    assert 'os.path.join(partition_root, "duplicates")' in source
    assert "_move_duplicate_folder_to_duplicates" in source
    assert "move_func(source, destination)" in source
    assert "send2trash" not in source


def test_build367_source_excludes_holding_folder_and_avoids_overwrite():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "exclude_paths=(holding_root,)" in source
    assert "_unique_duplicates_destination" in source
    assert "(moved {index})" in source


def test_build367_requirements_define_partition_root_move_behavior():
    text = _docx_text("TLO_Inventory_Requirements_Working_v373.docx")
    assert "folder named duplicates at the root of the partition" in text
    assert "shall create it before moving qualifying duplicates" in text
    assert "moved as one complete directory tree" in text
    assert "shall not overwrite" in text
    assert "exclude the duplicates holding folder from duplicate discovery" in text
    assert "Recycle Bin" not in text[text.find("18. Duplicate Cleanup Utility Requirements"):text.find("Appendix", text.find("18. Duplicate Cleanup Utility Requirements"))]


def test_build367_manual_documents_partition_root_move_behavior():
    text = (ROOT / "TLO_Inventory_User_Manual_v373.rtf").read_text(encoding="utf-8", errors="ignore")
    assert "folder named duplicates at the root of the partition" in text
    assert "creates that folder when it does not already exist" in text
    assert "moves the entire qualifying duplicate folder tree" in text
    assert "(moved 2)" in text


def test_build367_packaging_no_longer_collects_send2trash():
    for name in ("createWindowsDist.ps1", "createLinuxDist.sh", "createMacOSDist.sh"):
        text = (ROOT / name).read_text(encoding="utf-8")
        assert "tlo-deleteDupes.py" in text
        assert "imageio_ffmpeg" in text
        assert "send2trash" not in text


def test_build367_archives_build366_change_log():
    with ZipFile(ROOT / "old-change-logs.zip") as zf:
        assert "CHANGES_v366.txt" in zf.namelist()
