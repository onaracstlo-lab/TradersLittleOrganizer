"""Build 433 contract: GitHub Build Process remains outside the TLO source bundle."""
__version__ = "v433"

from pathlib import Path
from docx import Document
import pytest

pytestmark = pytest.mark.contract
ROOT = Path(__file__).resolve().parents[2]


def _doc_text(path: Path) -> str:
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            chunks.extend(cell.text for cell in row.cells)
    return "\n".join(chunks)


def test_build433_source_bundle_excludes_all_process_artifacts():
    names = {p.name for p in ROOT.iterdir() if p.is_file()}
    assert "Run-TLO-GitHub-Build.ps1" not in names
    assert "Create-TLOArtifactSigningMetadata.ps1" not in names
    assert not any(name.startswith("TLO_GitHub_Build_Process_Requirements_v") for name in names)
    assert not any(name.startswith("TLO_GitHub_Build_Process_v") and name.endswith(".zip") for name in names)


def test_build433_source_readme_documents_process_separation():
    text = (ROOT / "SOURCE_BUNDLE_README_v433.txt").read_text(encoding="utf-8")
    assert "No GitHub Build Process artifact is included in this source bundle." in text
    assert "independently versioned and distributed as its own separate package" in text


def test_build433_requirements_document_process_separation():
    text = _doc_text(ROOT / "TLO_Inventory_Requirements_Working_v433.docx")
    assert "shall contain no GitHub Build Process files" in text
    assert "TLO_GitHub_Build_Process_v<PROCESS>.zip" in text
    assert "transient build-process helpers" in text
