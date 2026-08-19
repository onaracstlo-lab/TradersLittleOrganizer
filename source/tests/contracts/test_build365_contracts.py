"""Build 375 FAQ and delete-duplicate contracts."""

__version__ = "v378"

from pathlib import Path

import pytest

pytestmark = pytest.mark.contract
ROOT = Path(__file__).resolve().parents[2]


def test_build365_faq_explains_first_music_folder_stops_descent():
    text = (ROOT / "TLO-FAQ.txt").read_text(encoding="utf-8")
    assert "Q: Does TLO inventory absolutely every music file?" in text
    assert "Once TLO identifies a music file, it stops looking." in text
    assert "nested folder with more music files" in text
    assert "opening act" in text


def test_build365_delete_dupes_moves_directory_path_as_one_object():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "_move_duplicate_folder_to_duplicates" in source
    assert "move_func(source, destination)" in source
    assert "duplicates holding folder" in source


def test_build365_delete_dupes_supports_same_artist_date_nonexact_names():
    source = (ROOT / "tlo-deleteDupes.py").read_text(encoding="utf-8")
    assert "_same_artist_and_date" in source
    assert "_potential_originals_for_copy" in source
    assert "same artist/date" in source.lower()


def _docx_text(name: str) -> str:
    from docx import Document
    doc = Document(ROOT / name)
    chunks = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    return "\n".join(chunks)


def test_build365_requirements_document_same_artist_date_and_folder_level_move():
    text = _docx_text("TLO_Inventory_Requirements_Working_v378.docx")
    assert "same-parent sibling directories" in text
    assert "normalized artist and date" in text
    assert "Candidate discovery alone shall never establish that two folders are duplicates" in text
    assert "moved as one complete directory tree" in text
    assert "duplicates" in text


def test_build365_manual_documents_same_artist_date_and_folder_level_move():
    text = (ROOT / "TLO_Inventory_User_Manual_v378.rtf").read_text(encoding="utf-8", errors="ignore")
    assert "normalized artist and date match" in text
    assert "A matching artist and date is only a way to select folders for comparison" in text
    assert "moved as a complete folder tree" in text
    assert "duplicates" in text


def test_build365_archives_build364_change_log():
    import zipfile
    with zipfile.ZipFile(ROOT / "old-change-logs.zip") as zf:
        assert "CHANGES_v364.txt" in zf.namelist()
